# -*- coding: utf-8 -*-
"""
assets_provider.py · S2 资料获取 AssetsProvider 抽象接口（契约定义）

本模块只定义 S2 阶段"资料获取"的抽象接口 + 一个占位实现,**不提供真实素材库 provider**:
- AssetsProvider: 抽象接口(lookup / resolve)。
- PlaceholderAssetsProvider: 占位实现(lookup 返回占位 / resolve 产占位 docx),仅供基建验证。

S2 的产物是 assets.json(手段无关结构 / 见 prompts/stages/s2-acquire.md);真实素材由调用方用其手段
(RAG / 检索 / 网络 / 人工)获取后写入 assets.json。本框架不内置检索引擎、不绑定特定素材库。

注(CuratedLocalAssetsProvider): 下方保留一个 CuratedLocalAssetsProvider 类作为"本地素材库 provider"的
接口示例,但它依赖未提供的映射模块、已在 get_provider() 中下架(get_provider('curated_local') 直接报错)。
不要把它当成可用的真实素材库——真实 provider 由你按 AssetsProvider 接口自行实现。

【未来实现(仅接口注释说明)】
MCPExternalAssetsProvider: 通过 MCP 协议对接外部材料库,适用于素材在统一平台集中管理的场景。
"""

from __future__ import annotations

from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
import tempfile


@dataclass
class AssetRef:
    """Assets 库中一个资源的引用。"""
    asset_type: str              # 资源类型,如 "资质证明" / "类似业绩" / "人员简历"
    is_placeholder: bool         # True = 占位引用,False = 真实命中
    lookup_key: str              # 查询键,便于未来查询溯源
    metadata: dict = field(default_factory=dict)  # 自由字段,供未来实现扩展


class AssetsProvider(ABC):
    """S2 资料获取 AssetsProvider 抽象接口。"""

    @abstractmethod
    def lookup(self, asset_type: str, **kwargs) -> AssetRef:
        """按 asset_type 和可选参数查找资源,返回 AssetRef。"""
        ...

    @abstractmethod
    def resolve(self, asset_ref: AssetRef) -> Path:
        """把 AssetRef 解析为实际可用文件路径(docx / pdf 等)。"""
        ...


class PlaceholderAssetsProvider(AssetsProvider):
    """所有 lookup 返回占位 AssetRef,resolve 产出占位 docx。

    不查任何真实资源。仅供接口基建验证 / 占位运行。
    """

    def lookup(self, asset_type: str, **kwargs) -> AssetRef:
        # 构造查询键(便于溯源,即使本轮不用)
        kv_parts = [f"{k}={v}" for k, v in sorted(kwargs.items())]
        lookup_key = f"{asset_type}|{'|'.join(kv_parts)}" if kv_parts else asset_type
        return AssetRef(
            asset_type=asset_type,
            is_placeholder=True,
            lookup_key=lookup_key,
            metadata={"reason": "PlaceholderAssetsProvider: 本轮不做真实查找"},
        )

    def resolve(self, asset_ref: AssetRef) -> Path:
        """生成含占位文字的临时 docx 段落文件,存放于临时目录(不进 tracked_outputs)。"""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("缺少 python-docx 依赖")

        doc = Document()
        doc.add_paragraph(f"[此处插入 {asset_ref.asset_type} 材料]")
        doc.add_paragraph(f"  lookup_key: {asset_ref.lookup_key}")
        doc.add_paragraph("  (由 PlaceholderAssetsProvider 产出,真实材料待填)")

        # 临时目录,不进基线
        tmp_dir = Path(tempfile.gettempdir()) / "solution_drafter_placeholder_assets"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        # 用 lookup_key 的哈希做文件名避免冲突
        import hashlib
        fname = f"placeholder_{hashlib.md5(asset_ref.lookup_key.encode('utf-8')).hexdigest()[:12]}.docx"
        out_path = tmp_dir / fname
        doc.save(str(out_path))
        return out_path


class CuratedLocalAssetsProvider(AssetsProvider):
    """本地素材库 provider 的接口示例（已在 get_provider 中下架 / 依赖未提供的映射模块 / 勿当真实素材库用）。

    扫 assets/<类别>/<company_id>/ 结构找真实 asset 文件。

    候选优先级:_raw/<file>.docx → _raw/<file>.pdf(走 pdf2docx 转换) →
    <resource>.md(curated 兜底,转 docx 段落)。

    多命中:stdin 是 tty → 让用户选;非 tty 或 non_interactive=True →
    按 lookup_priority 自动选(默认 latest_year_first)。

    asset_type → 类别 的映射查 references/asset_type_mapping.yaml。
    """

    def __init__(
        self,
        assets_root: Path | None = None,
        company_id: str = "own_default",
        non_interactive: bool = False,
        mapping_path: Path | None = None,
    ):
        self.assets_root = (
            Path(assets_root) if assets_root
            else Path(__file__).resolve().parent.parent / "assets"
        )
        self.company_id = company_id
        self.non_interactive = non_interactive
        self.mapping_path = mapping_path

    # ─── 内部:扫描候选 ───

    def _scan_candidates(
        self,
        asset_type: str,
        asset_query_type: str = "",
        year_filter: str | None = None,
    ) -> list[AssetRef]:
        # 延迟 import,避免顶层依赖
        # _asset_type_mapping 是本类依赖的映射模块 / 本框架未提供 (CuratedLocal 已下架)
        # 仅 CuratedLocalAssetsProvider 的 lookup 路径依赖 / PlaceholderAssetsProvider 不受影响
        import sys as _sys
        _sys.path.insert(0, str(Path(__file__).resolve().parent))
        try:
            from _asset_type_mapping import (
                map_asset_type_to_category,
                PLACEHOLDER_SENTINEL,
            )
        except ImportError as exc:
            raise RuntimeError(
                "CuratedLocalAssetsProvider 依赖 _asset_type_mapping 未迁移 "
                "(本框架未提供该映射模块 / CuratedLocal 已下架)。请用 PlaceholderAssetsProvider 或自行实现 AssetsProvider "
                "(get_provider('placeholder'))。"
            ) from exc

        category = map_asset_type_to_category(
            asset_type, asset_query_type, mapping_path=self.mapping_path,
        )
        if category == PLACEHOLDER_SENTINEL:
            return []

        category_dir = self.assets_root / category / self.company_id
        if not category_dir.exists():
            return []

        candidates: list[AssetRef] = []

        # _raw/ 优先 docx,然后 pdf
        raw_dir = category_dir / "_raw"
        if raw_dir.exists():
            for path in sorted(raw_dir.glob("*.docx")):
                candidates.append(self._make_ref(path, asset_type, "raw_docx"))
            for path in sorted(raw_dir.glob("*.pdf")):
                candidates.append(self._make_ref(path, asset_type, "raw_pdf"))

        # 仅当 _raw 全空时,curated md 兜底
        if not candidates:
            for path in sorted(category_dir.glob("*.md")):
                stem = path.stem
                # 跳过索引/schema/README 类
                if stem in ("README", "资质清单", "业绩列表", "简历索引") \
                        or stem.endswith("schema"):
                    continue
                candidates.append(self._make_ref(path, asset_type, "curated_md"))

        # year_filter
        if year_filter:
            candidates = [c for c in candidates if self._year_matches(c, year_filter)]

        return candidates

    @staticmethod
    def _make_ref(source_path: Path, asset_type: str, kind: str) -> AssetRef:
        """从文件名提取 year(若文件名以 YYYYMMDD_ 起头)。"""
        import re
        year = None
        m = re.match(r"^(\d{4})\d{2}\d{2}_", source_path.name)
        if m:
            year = int(m.group(1))
        return AssetRef(
            asset_type=asset_type,
            is_placeholder=False,
            lookup_key=f"{asset_type}|{source_path.name}",
            metadata={
                "source_path": str(source_path),
                "kind": kind,
                "year": year,
                "filename": source_path.name,
            },
        )

    @staticmethod
    def _year_matches(ref: AssetRef, year_filter: str) -> bool:
        year = ref.metadata.get("year")
        if year is None:
            return False
        try:
            y = int(year)
        except (ValueError, TypeError):
            return False
        s = year_filter.strip()
        if s.startswith(">="):
            return y >= int(s[2:])
        if s.startswith("<="):
            return y <= int(s[2:])
        if "-" in s:
            lo, hi = s.split("-", 1)
            return int(lo) <= y <= int(hi)
        return y == int(s)

    # ─── lookup ───

    def lookup(self, asset_type: str, **kwargs) -> AssetRef:
        asset_query_type = kwargs.get("asset_query_type", "")
        year_filter = kwargs.get("year_filter", None)
        lookup_priority = kwargs.get("lookup_priority", "latest_year_first")

        candidates = self._scan_candidates(
            asset_type,
            asset_query_type=asset_query_type,
            year_filter=year_filter,
        )

        if len(candidates) == 0:
            return AssetRef(
                asset_type=asset_type,
                is_placeholder=True,
                lookup_key=f"{asset_type}|miss",
                metadata={
                    "reason": (
                        f"CuratedLocalAssetsProvider: 0 命中 in "
                        f"{self.company_id} (asset_type={asset_type!r})"
                    ),
                },
            )

        if len(candidates) == 1:
            return candidates[0]

        return self._choose(candidates, asset_type, lookup_priority)

    def _choose(
        self, candidates: list[AssetRef], asset_type: str, lookup_priority: str,
    ) -> AssetRef:
        """多命中选择:tty 走交互,非 tty 按 priority 自动。"""
        import sys as _sys
        if self.non_interactive or not _sys.stdin.isatty():
            return self._sort_by_priority(candidates, lookup_priority)[0]

        # tty 交互
        print(
            f"\n[lookup 多命中] asset_type='{asset_type}' 共 {len(candidates)} 个候选:",
            file=_sys.stderr,
        )
        for i, c in enumerate(candidates, 1):
            print(
                f"  [{i}] {c.metadata.get('filename')} "
                f"(kind={c.metadata.get('kind')}, year={c.metadata.get('year')})",
                file=_sys.stderr,
            )
        while True:
            print(f"  请输入 1-{len(candidates)}: ", end="", file=_sys.stderr, flush=True)
            try:
                line = _sys.stdin.readline().strip()
                idx = int(line)
                if 1 <= idx <= len(candidates):
                    return candidates[idx - 1]
            except (ValueError, EOFError):
                pass

    @staticmethod
    def _sort_by_priority(
        candidates: list[AssetRef], lookup_priority: str,
    ) -> list[AssetRef]:
        # 本期实现 latest_year_first;其他 priority 留 V4,fallback 到 latest_year_first
        return sorted(
            candidates,
            key=lambda c: -(c.metadata.get("year") or 0),
        )

    # ─── resolve ───

    def resolve(self, asset_ref: AssetRef) -> Path:
        if asset_ref.is_placeholder:
            return PlaceholderAssetsProvider().resolve(asset_ref)

        source_path = Path(asset_ref.metadata["source_path"])
        kind = asset_ref.metadata.get("kind", "")

        if kind == "raw_docx":
            return source_path
        if kind == "raw_pdf":
            return self._pdf_to_docx(source_path, asset_ref)
        if kind == "curated_md":
            return self._md_to_docx(source_path, asset_ref)
        return PlaceholderAssetsProvider().resolve(asset_ref)

    @staticmethod
    def _pdf_to_docx(pdf_path: Path, asset_ref: AssetRef) -> Path:
        import sys as _sys
        try:
            from pdf2docx import Converter
        except ImportError:
            print(
                f"[警告] pdf2docx 未安装,降级 placeholder for {pdf_path.name}",
                file=_sys.stderr,
            )
            return PlaceholderAssetsProvider().resolve(asset_ref)

        out = Path(tempfile.gettempdir()) / "solution_drafter_pdf2docx" / f"{pdf_path.stem}.docx"
        out.parent.mkdir(parents=True, exist_ok=True)
        try:
            cv = Converter(str(pdf_path))
            cv.convert(str(out))
            cv.close()
            return out
        except Exception as exc:
            print(
                f"[警告] pdf2docx 转换失败 {pdf_path.name}: {exc!r},降级 placeholder",
                file=_sys.stderr,
            )
            return PlaceholderAssetsProvider().resolve(asset_ref)

    @staticmethod
    def _md_to_docx(md_path: Path, asset_ref: AssetRef) -> Path:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("缺少 python-docx 依赖")
        text = md_path.read_text(encoding="utf-8")
        doc = Document()
        doc.add_paragraph(f"[curated md] {md_path.stem}")
        for line in text.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)
        out = Path(tempfile.gettempdir()) / "solution_drafter_md_assets" / f"{md_path.stem}.docx"
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return out


def get_provider(name: str, **kwargs) -> AssetsProvider:
    """按 manifest.yaml 的 assets_provider 字段返回对应实例。

    kwargs 透传给 Provider 构造器(用于 CuratedLocalAssetsProvider 的 company_id 等)。
    """
    if name == "placeholder":
        return PlaceholderAssetsProvider()
    if name == "curated_local":
        # 下架:CuratedLocalAssetsProvider 依赖本框架未提供的映射模块 + pdf2docx(第三方未引入)/
        # 不从 get_provider 公开。类定义保留作接口示例。
        raise ValueError(
            f"assets_provider='curated_local' 已下架(依赖未提供的映射模块 / pdf2docx 第三方未引入)。"
            f"当前支持: 'placeholder' / 或按 AssetsProvider 接口自行实现。"
        )
    raise ValueError(
        f"未知的 assets_provider='{name}'。当前支持: 'placeholder';"
        f"未来的 MCPExternalAssetsProvider 见本模块文档。"
    )
