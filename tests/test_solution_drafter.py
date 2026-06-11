# -*- coding: utf-8 -*-
"""
test_solution_drafter.py · 端到端 + 契约回归

覆盖:
- 3 个文档类型 outline.yaml 通过 schema 校验
- 4 个 demo 端到端跑通 (draft 二级标题数 = section 数 / final.docx 落盘 / 字体合规)
- 逐 section append+save 落盘 → docx 含全部 section
- 模拟中途断: 写 2 section + save → 重新读回 → 2 个保住在磁盘 (长文档不断流核心)
- 领域插件双示例对照: 默认空领域 §公司实力 = 【待补充】 vs 虚构领域 = own_* 填充
"""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

import pytest
import yaml
from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT / "scripts"))
from docx_builder import create_section_doc
from append_chapter import append_markdown
from check_font_safety import check_font_safety
from validate_outline import validate
from llm_stub import _field, PLACEHOLDER
sys.path.insert(0, str(REPO_ROOT / "examples"))
from _demo_s5 import s5_review_and_render

DOC_TYPES = ["需求方案", "解决方案建议书", "实施方案"]
DEMOS = ["demo-需求方案", "demo-解决方案建议书", "demo-实施方案", "demo-槐序数据领域"]


def _heading_texts(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]


# ── outline 校验 ──────────────────────────────────────────────

@pytest.mark.parametrize("doc_type", DOC_TYPES)
def test_outline_validates(doc_type):
    yaml_path = REPO_ROOT / "templates" / doc_type / "outline.yaml"
    result = subprocess.run(
        [sys.executable, "scripts/validate_outline.py", str(yaml_path), "--repo-root", "."],
        cwd=str(REPO_ROOT), capture_output=True, text=True,
    )
    assert result.returncode == 0, f"{doc_type} validate FAIL: {result.stdout} {result.stderr}"


@pytest.mark.parametrize("doc_type", DOC_TYPES)
def test_outline_sections_have_asset_needs(doc_type):
    """每 section 必有非空 asset_needs (S2 素材链)."""
    cfg = yaml.safe_load((REPO_ROOT / "templates" / doc_type / "outline.yaml").read_text(encoding="utf-8"))
    for sec in cfg["outline"]:
        assert sec.get("asset_needs"), f"{doc_type}/{sec['id']} 缺 asset_needs"


# ── demo 端到端 ───────────────────────────────────────────────

@pytest.mark.parametrize("demo", DEMOS)
def test_demo_end_to_end(demo):
    demo_dir = REPO_ROOT / "examples" / demo
    out_dir = demo_dir / "output"
    if out_dir.exists():
        shutil.rmtree(out_dir)
    result = subprocess.run(
        [sys.executable, f"examples/{demo}/run.py"],
        cwd=str(REPO_ROOT), capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    assert result.returncode == 0, f"{demo} demo FAIL: {result.stderr}"
    # draft 二级标题数 = outline section 数
    draft = (out_dir / "draft.md").read_text(encoding="utf-8")
    assert draft.count("\n## ") == 6, f"{demo} draft 二级标题应 6 / 实测 {draft.count(chr(10) + '## ')}"
    # final.docx 落盘 + 字体合规
    docx_files = list(out_dir.glob("*.docx"))
    assert len(docx_files) == 1, f"{demo} final.docx 未唯一落盘"
    assert docx_files[0].stat().st_size > 10000
    assert check_font_safety(docx_files[0]) == [], f"{demo} 字体不合规"
    # S5 真闭环: final.md 产出 (不再空壳) + 自检透传 (== draft)
    assert (out_dir / "final.md").exists(), f"{demo} S5 未产 final.md (空壳回归)"
    assert (out_dir / "final.md").read_text(encoding="utf-8") == draft, f"{demo} final.md 应自检透传 == draft"
    # ★ S2→S4 素材消费 (验真正价值 / 不只验结构): draft 真含 acquired 素材 + 待补充落【待补充】
    assert "依据已获取的" in draft, f"{demo} draft 未见素材消费句 (assets 未被消费)"
    assert "示例素材" in draft, f"{demo} draft 未含 S2 acquired 素材内容"
    assert "【待补充】" in draft, f"{demo} draft 未含【待补充】 (待补充素材未显性化)"


# ── 长文档不断流 (逐段落盘) ───────────────────────────────────

def test_incremental_append_docx_complete(tmp_path):
    """逐 section append+save → docx 含全部 section."""
    docx_path = tmp_path / "t.docx"
    create_section_doc(docx_path, body_font="宋体")
    document = Document(str(docx_path))
    for i in range(1, 6):
        append_markdown(document, f"## 第{i}节\n\n正文{i}。\n", body_font="宋体")
        document.save(str(docx_path))   # 逐段落盘
    titles = [h for h in _heading_texts(docx_path) if h.startswith("第")]
    assert len(titles) == 5, f"应 5 个 section / 实测 {titles}"


def test_midway_break_keeps_written_sections(tmp_path):
    """模拟中途断: append 2 section+save → 新 Document 读回 → 2 个在磁盘 (逐段落盘真保住)."""
    docx_path = tmp_path / "t.docx"
    create_section_doc(docx_path, body_font="宋体")
    document = Document(str(docx_path))
    append_markdown(document, "## 甲节\n\n甲内容。\n", body_font="宋体")
    document.save(str(docx_path))
    append_markdown(document, "## 乙节\n\n乙内容。\n", body_font="宋体")
    document.save(str(docx_path))
    del document   # 模拟进程结束 / 内存态丢失
    titles = [h for h in _heading_texts(docx_path) if h in ("甲节", "乙节")]
    assert "甲节" in titles and "乙节" in titles and len(titles) == 2


# ── 领域插件双示例对照 ────────────────────────────────────────

def test_domain_plugin_contrast():
    """默认空领域 §公司实力 = 【待补充】 vs 虚构领域 = own_* 填充 (L3 接入效果)."""
    # 两 demo 已由 test_demo_end_to_end 跑过 (output 存在)
    empty = (REPO_ROOT / "examples" / "demo-解决方案建议书" / "output" / "draft.md")
    domain = (REPO_ROOT / "examples" / "demo-槐序数据领域" / "output" / "draft.md")
    if not empty.exists():
        subprocess.run([sys.executable, "examples/demo-解决方案建议书/run.py"],
                       cwd=str(REPO_ROOT), capture_output=True)
    if not domain.exists():
        subprocess.run([sys.executable, "examples/demo-槐序数据领域/run.py"],
                       cwd=str(REPO_ROOT), capture_output=True)
    empty_txt = empty.read_text(encoding="utf-8")
    domain_txt = domain.read_text(encoding="utf-8")
    # 空领域: 公司实力段含【待补充】 / 不含虚构公司名
    company_block = empty_txt.split("## 公司实力", 1)[1].split("\n## ", 1)[0]
    assert "【待补充】" in company_block, "空领域 §公司实力 应为【待补充】"
    assert "槐序数据" not in empty_txt, "空领域不应出现领域公司名"
    # 虚构领域: 公司实力段由 own_* 填充 (含虚构公司名)
    assert "槐序数据科技" in domain_txt, "虚构领域 §公司实力 应由 own_* 填充"


# ── S5 真闭环 (draft → s5-review 自检 → final.md + final.docx) ──────────

def test_s5_helper_produces_final_and_consumes_review(tmp_path):
    """S5 helper: 自检透传 → final.md 产出 + final.docx 渲染 + s5-review 非空被消费."""
    outline_cfg = {"outline": [{"id": "01"}, {"id": "02"}], "output": {"font_policy": "宋体"}}
    doc_title = "# 测试项目 需求方案\n"
    sections = ["## 甲节\n\n甲内容。\n", "## 乙节\n\n乙内容。\n"]
    draft = doc_title + "\n" + "\n".join(sections)
    final_md = tmp_path / "final.md"
    final_docx = tmp_path / "out.docx"
    sc, font, stats = s5_review_and_render(
        draft_text=draft, section_markdowns=sections, doc_title=doc_title,
        outline_cfg=outline_cfg, s5_review_md="（s5-review 清单字面）",
        final_md_path=final_md, final_docx_path=final_docx,
    )
    assert final_md.exists() and final_md.read_text(encoding="utf-8") == draft
    heads = [p.text for p in Document(str(final_docx)).paragraphs if p.style.name.startswith("Heading")]
    assert "甲节" in heads and "乙节" in heads
    assert sc == [] and font == []


def test_s5_self_check_catches_leak():
    """S5 自检: section 数不符 + 内部术语外泄 → 报问题 (不静默)."""
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        sc, _, _ = s5_review_and_render(
            draft_text="# T\n\n## 甲\n\nintake_fields 泄漏\n",
            section_markdowns=["## 甲\n\nintake_fields 泄漏\n"], doc_title="# T\n",
            outline_cfg={"outline": [{"id": "01"}, {"id": "02"}], "output": {"font_policy": "宋体"}},
            s5_review_md="x", final_md_path=td / "f.md", final_docx_path=td / "o.docx",
        )
    assert any("二级标题数" in i for i in sc) and any("intake_fields" in i for i in sc)


def test_s5_empty_review_rejected(tmp_path):
    """s5-review.md 为空 → 断言失败 (S5 stage prompt 必须加载 / 证被消费)."""
    with pytest.raises(AssertionError):
        s5_review_and_render(
            draft_text="# T\n\n## 甲\n\nx\n", section_markdowns=["## 甲\n\nx\n"], doc_title="# T\n",
            outline_cfg={"outline": [{"id": "01"}], "output": {"font_policy": "宋体"}}, s5_review_md="",
            final_md_path=tmp_path / "f.md", final_docx_path=tmp_path / "o.docx",
        )


# ── 契约校验 (font_policy 白名单 / intake 缺失守) ──────────────────────

def test_font_policy_must_be_whitelisted(tmp_path):
    """font_policy 不在 check_font_safety 白名单 → validate 拦 (避免渲染期才炸)."""
    src = (REPO_ROOT / "templates" / "需求方案" / "outline.yaml").read_text(encoding="utf-8")
    bad = tmp_path / "bad.yaml"
    bad.write_text(src.replace("font_policy: 宋体", "font_policy: 楷体"), encoding="utf-8")
    assert any("font_policy" in i and "白名单" in i for i in validate(bad, REPO_ROOT)), "楷体 应被拦"
    good = tmp_path / "good.yaml"
    good.write_text(src, encoding="utf-8")
    assert not any("font_policy" in i for i in validate(good, REPO_ROOT)), "宋体 不应报"


def test_missing_required_field_to_placeholder():
    """intake 缺 required 字段 → 【待补充】 (不静默 / 不编造)."""
    assert _field({}, "project_name") == PLACEHOLDER
    assert _field({"project_name": None}, "project_name") == PLACEHOLDER
    assert _field({"project_name": "X项目"}, "project_name") == "X项目"
